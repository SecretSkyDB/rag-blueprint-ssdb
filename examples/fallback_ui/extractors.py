"""Document text extractors for the fallback-ui upload endpoint.

One function per format, all returning ``(text, meta)`` where ``meta`` is a
dict with at least ``{"chars": int, "format": str}``. Anything that fails
falls back to ``extract_text_best_effort()`` which tries utf-8 then
latin-1 — same as the original ``app.upload()`` did for arbitrary bytes.

Kept dependency-light on purpose: every import is wrapped so the module
loads even if optional extras are missing. The ``upload()`` handler then
returns a clean error message naming the missing extra.
"""
from __future__ import annotations

import io
import re
from typing import Callable

# ── Optional imports (each guarded so app boots if any wheel is absent) ───
try:
    from pypdf import PdfReader  # type: ignore
except Exception:  # pragma: no cover - import-time only
    PdfReader = None  # type: ignore

try:
    import docx  # type: ignore  # python-docx
except Exception:  # pragma: no cover
    docx = None  # type: ignore

try:
    from bs4 import BeautifulSoup  # type: ignore
except Exception:  # pragma: no cover
    BeautifulSoup = None  # type: ignore

try:
    from ebooklib import epub, ITEM_DOCUMENT  # type: ignore
except Exception:  # pragma: no cover
    epub = None  # type: ignore
    ITEM_DOCUMENT = None  # type: ignore

try:
    from striprtf.striprtf import rtf_to_text  # type: ignore
except Exception:  # pragma: no cover
    rtf_to_text = None  # type: ignore


# ── Per-format extractors ─────────────────────────────────────────────────


def _extract_pdf(raw: bytes) -> str:
    if PdfReader is None:
        raise RuntimeError("pypdf is not installed (pip install pypdf)")
    reader = PdfReader(io.BytesIO(raw))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(p for p in parts if p.strip())


def _extract_docx(raw: bytes) -> str:
    if docx is None:
        raise RuntimeError("python-docx is not installed (pip install python-docx)")
    d = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in d.paragraphs if p.text]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_html(raw: bytes) -> str:
    if BeautifulSoup is None:
        raise RuntimeError("beautifulsoup4 is not installed (pip install beautifulsoup4 lxml)")
    try:
        soup = BeautifulSoup(raw, "lxml")
    except Exception:
        soup = BeautifulSoup(raw, "html.parser")
    for bad in soup(["script", "style", "noscript"]):
        bad.decompose()
    text = soup.get_text(" ")
    return re.sub(r"[ \t]+", " ", re.sub(r"\n{3,}", "\n\n", text)).strip()


def _extract_epub(raw: bytes) -> str:
    if epub is None:
        raise RuntimeError("EbookLib is not installed (pip install EbookLib)")
    with io.BytesIO(raw) as buf:
        # ebooklib reads from disk; round-trip via tempfile if needed.
        import tempfile
        import os as _os
        with tempfile.NamedTemporaryFile(delete=False, suffix=".epub") as tf:
            tf.write(buf.read())
            tmp_path = tf.name
        try:
            book = epub.read_epub(tmp_path)
        finally:
            try:
                _os.unlink(tmp_path)
            except OSError:
                pass
    parts: list[str] = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        html = item.get_content()
        if BeautifulSoup is not None:
            try:
                soup = BeautifulSoup(html, "lxml")
            except Exception:
                soup = BeautifulSoup(html, "html.parser")
            for bad in soup(["script", "style"]):
                bad.decompose()
            parts.append(soup.get_text(" "))
        else:
            parts.append(html.decode("utf-8", errors="replace"))
    text = "\n\n".join(parts)
    return re.sub(r"[ \t]+", " ", re.sub(r"\n{3,}", "\n\n", text)).strip()


def _extract_rtf(raw: bytes) -> str:
    if rtf_to_text is None:
        raise RuntimeError("striprtf is not installed (pip install striprtf)")
    try:
        s = raw.decode("utf-8")
    except UnicodeDecodeError:
        s = raw.decode("latin-1", errors="replace")
    return rtf_to_text(s)


def _extract_plain(raw: bytes) -> str:
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1", errors="replace")


# ── Dispatcher ────────────────────────────────────────────────────────────


_EXT_MAP: dict[str, Callable[[bytes], str]] = {
    "pdf":   _extract_pdf,
    "docx":  _extract_docx,
    "html":  _extract_html,
    "htm":   _extract_html,
    "xhtml": _extract_html,
    "epub":  _extract_epub,
    "rtf":   _extract_rtf,
    # plain-text family — handled by _extract_plain
    "txt":   _extract_plain,
    "md":    _extract_plain,
    "markdown": _extract_plain,
    "rst":   _extract_plain,
    "csv":   _extract_plain,
    "json":  _extract_plain,
    "log":   _extract_plain,
    "py":    _extract_plain,
    "yml":   _extract_plain,
    "yaml":  _extract_plain,
}


_MIME_MAP: dict[str, Callable[[bytes], str]] = {
    "application/pdf":                                                          _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document":  _extract_docx,
    "application/msword":                                                       _extract_docx,
    "text/html":                                                                _extract_html,
    "application/xhtml+xml":                                                    _extract_html,
    "application/epub+zip":                                                     _extract_epub,
    "application/rtf":                                                          _extract_rtf,
    "text/rtf":                                                                 _extract_rtf,
}


def extract(raw: bytes, *, filename: str = "", mimetype: str = "") -> tuple[str, dict]:
    """Extract text from ``raw`` bytes using filename/mimetype hints.

    Returns ``(text, meta)``. ``meta`` includes ``{"format": "...", "chars": N}``.
    Unknown formats fall through to plain decoding (utf-8 then latin-1) and
    report ``format="plain"`` so callers can warn the user that ingestion
    is best-effort.
    """
    fn = (filename or "").lower()
    mt = (mimetype or "").lower().split(";", 1)[0].strip()

    fn_ext = fn.rsplit(".", 1)[-1] if "." in fn else ""

    extractor = _MIME_MAP.get(mt) or _EXT_MAP.get(fn_ext) or _extract_plain
    fmt = (
        "pdf"  if extractor is _extract_pdf  else
        "docx" if extractor is _extract_docx else
        "html" if extractor is _extract_html else
        "epub" if extractor is _extract_epub else
        "rtf"  if extractor is _extract_rtf  else
        "plain"
    )
    text = extractor(raw)
    return text, {"format": fmt, "chars": len(text), "bytes_in": len(raw)}


__all__ = ["extract"]
